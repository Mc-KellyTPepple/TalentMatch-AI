"""
TalentMatch AI
Resume Parser

Responsibilities:
- Parse PDF resumes using pypdf
- Parse DOCX resumes using python-docx
- Parse TXT resumes
- Validate file types
- Clean extracted text
- Provide useful errors
- Designed for Render Free / 512 MB RAM

Supported:
- PDF
- DOCX
- TXT

PDF engine:
- pypdf

DOCX engine:
- python-docx

IMPORTANT:
This file does NOT use PyMuPDF / fitz.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from pypdf import PdfReader
from docx import Document


# ============================================================
# Custom Exception
# ============================================================

class ResumeParsingError(Exception):
    """
    Raised when a resume cannot be parsed.
    """
    pass


# ============================================================
# Supported Extensions
# ============================================================

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
}


# ============================================================
# Text Cleaning
# ============================================================

def clean_text(text: str) -> str:
    """
    Clean extracted resume text.

    Removes:
    - Null characters
    - Excess whitespace
    - Repeated blank lines
    - Unnecessary spacing

    Returns:
        Clean readable text.
    """

    if not text:
        return ""

    # Remove null characters
    text = text.replace("\x00", " ")

    # Normalize line endings
    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")

    # Normalize non-breaking spaces
    text = text.replace("\xa0", " ")

    # Normalize spaces and tabs
    text = re.sub(
        r"[ \t]+",
        " ",
        text,
    )

    # Remove excessive blank lines
    text = re.sub(
        r"\n\s*\n\s*\n+",
        "\n\n",
        text,
    )

    # Clean individual lines
    lines = []

    for line in text.split("\n"):
        line = line.strip()

        if line:
            lines.append(line)

    text = "\n".join(lines)

    return text.strip()


# ============================================================
# PDF Parser
# ============================================================

def parse_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF using pypdf.

    No PyMuPDF dependency is used.

    Parameters:
        file_bytes:
            Raw PDF bytes.

    Returns:
        Extracted and cleaned text.

    Raises:
        ResumeParsingError
    """

    if not file_bytes:
        raise ResumeParsingError(
            "The PDF file is empty."
        )

    pdf_stream = None

    try:

        # ----------------------------------------------------
        # Basic PDF signature validation
        # ----------------------------------------------------

        if not file_bytes.startswith(b"%PDF"):
            raise ResumeParsingError(
                "The uploaded file is not a valid PDF. "
                "Please upload a genuine PDF resume."
            )

        # ----------------------------------------------------
        # Load PDF from memory
        # ----------------------------------------------------

        pdf_stream = io.BytesIO(file_bytes)

        # ----------------------------------------------------
        # IMPORTANT:
        #
        # This uses pypdf.
        #
        # There is deliberately NO:
        #
        # import fitz
        # fitz.open(...)
        #
        # ----------------------------------------------------

        reader = PdfReader(pdf_stream)

        # ----------------------------------------------------
        # Check encrypted PDF
        # ----------------------------------------------------

        if reader.is_encrypted:
            raise ResumeParsingError(
                "This PDF is password protected. "
                "Please upload an unlocked PDF resume."
            )

        # ----------------------------------------------------
        # Check pages
        # ----------------------------------------------------

        if len(reader.pages) == 0:
            raise ResumeParsingError(
                "The PDF does not contain any pages."
            )

        extracted_pages = []

        # ----------------------------------------------------
        # Extract text from every page
        # ----------------------------------------------------

        for page_number, page in enumerate(
            reader.pages,
            start=1,
        ):

            try:

                page_text = page.extract_text()

                if page_text:

                    page_text = clean_text(
                        page_text
                    )

                    if page_text:
                        extracted_pages.append(
                            page_text
                        )

            except Exception as exc:

                print(
                    f"PDF page {page_number} "
                    f"text extraction warning: "
                    f"{repr(exc)}"
                )

                # Continue with other pages
                continue

        # ----------------------------------------------------
        # Combine pages
        # ----------------------------------------------------

        text = "\n\n".join(
            extracted_pages
        )

        text = clean_text(
            text
        )

        # ----------------------------------------------------
        # No text extracted
        # ----------------------------------------------------

        if not text:

            raise ResumeParsingError(
                "No readable text could be extracted "
                "from this PDF. The PDF may be scanned "
                "or image-based. Please upload a text-based "
                "PDF, DOCX or TXT resume."
            )

        return text

    except ResumeParsingError:
        raise

    except Exception as exc:

        print(
            "PDF parsing error:",
            repr(exc),
        )

        raise ResumeParsingError(
            "Unable to read the PDF resume. "
            "Please make sure the PDF is valid, "
            "not corrupted and not password protected."
        ) from exc

    finally:

        if pdf_stream is not None:

            try:
                pdf_stream.close()

            except Exception:
                pass


# ============================================================
# DOCX Parser
# ============================================================

def parse_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX resume.

    Includes:
    - Paragraphs
    - Tables

    Parameters:
        file_bytes:
            Raw DOCX bytes.

    Returns:
        Extracted and cleaned text.

    Raises:
        ResumeParsingError
    """

    if not file_bytes:

        raise ResumeParsingError(
            "The DOCX file is empty."
        )

    document_stream = None

    try:

        # ----------------------------------------------------
        # Load DOCX from memory
        # ----------------------------------------------------

        document_stream = io.BytesIO(
            file_bytes
        )

        document = Document(
            document_stream
        )

        parts = []

        # ----------------------------------------------------
        # Paragraphs
        # ----------------------------------------------------

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:

                parts.append(
                    text
                )

        # ----------------------------------------------------
        # Tables
        # ----------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                cells = []

                for cell in row.cells:

                    cell_text = cell.text.strip()

                    if cell_text:

                        cells.append(
                            cell_text
                        )

                if cells:

                    parts.append(
                        " | ".join(cells)
                    )

        # ----------------------------------------------------
        # Combine and clean
        # ----------------------------------------------------

        text = "\n".join(
            parts
        )

        text = clean_text(
            text
        )

        # ----------------------------------------------------
        # Check result
        # ----------------------------------------------------

        if not text:

            raise ResumeParsingError(
                "No readable text could be extracted "
                "from this DOCX resume."
            )

        return text

    except ResumeParsingError:
        raise

    except Exception as exc:

        print(
            "DOCX parsing error:",
            repr(exc),
        )

        raise ResumeParsingError(
            "Unable to read the DOCX resume. "
            "Please make sure the document is valid "
            "and not corrupted."
        ) from exc

    finally:

        if document_stream is not None:

            try:
                document_stream.close()

            except Exception:
                pass


# ============================================================
# TXT Parser
# ============================================================

def parse_txt(file_bytes: bytes) -> str:
    """
    Extract text from a TXT resume.

    Attempts UTF-8 first, then common fallbacks.
    """

    if not file_bytes:

        raise ResumeParsingError(
            "The TXT file is empty."
        )

    encodings = [
        "utf-8",
        "utf-8-sig",
        "cp1252",
        "latin-1",
    ]

    text = None

    for encoding in encodings:

        try:

            text = file_bytes.decode(
                encoding
            )

            break

        except UnicodeDecodeError:

            continue

    if text is None:

        raise ResumeParsingError(
            "Unable to decode the TXT resume."
        )

    text = clean_text(
        text
    )

    if not text:

        raise ResumeParsingError(
            "The TXT resume contains no readable text."
        )

    return text


# ============================================================
# Main Resume Parser
# ============================================================

def parse_resume(
    file_bytes: bytes,
    filename: str,
) -> str:
    """
    Automatically determine the resume format
    and extract its text.

    Parameters:
        file_bytes:
            Uploaded resume bytes.

        filename:
            Original uploaded filename.

    Returns:
        Clean resume text.

    Raises:
        ResumeParsingError
    """

    # --------------------------------------------------------
    # Validate bytes
    # --------------------------------------------------------

    if not file_bytes:

        raise ResumeParsingError(
            "The uploaded resume is empty."
        )

    # --------------------------------------------------------
    # Validate filename
    # --------------------------------------------------------

    if not filename:

        raise ResumeParsingError(
            "The uploaded file has no filename."
        )

    # --------------------------------------------------------
    # Determine extension
    # --------------------------------------------------------

    extension = (
        Path(filename)
        .suffix
        .lower()
    )

    # --------------------------------------------------------
    # Validate extension
    # --------------------------------------------------------

    if extension not in SUPPORTED_EXTENSIONS:

        raise ResumeParsingError(
            "Unsupported resume format. "
            "Please upload a PDF, DOCX or TXT file."
        )

    # --------------------------------------------------------
    # Route PDF
    # --------------------------------------------------------

    if extension == ".pdf":

        return parse_pdf(
            file_bytes
        )

    # --------------------------------------------------------
    # Route DOCX
    # --------------------------------------------------------

    if extension == ".docx":

        return parse_docx(
            file_bytes
        )

    # --------------------------------------------------------
    # Route TXT
    # --------------------------------------------------------

    if extension == ".txt":

        return parse_txt(
            file_bytes
        )

    # --------------------------------------------------------
    # Safety fallback
    # --------------------------------------------------------

    raise ResumeParsingError(
        "Unable to determine the resume file format."
    )


# ============================================================
# Parser Status
# ============================================================

def parser_status():
    """
    Return parser availability information.

    Useful for diagnostics.
    """

    return {
        "status": "ready",
        "pdf_engine": "pypdf",
        "docx_engine": "python-docx",
        "txt_engine": "built-in Python",
        "pymupdf_required": False,
        "supported_formats": [
            "PDF",
            "DOCX",
            "TXT",
        ],
    }
